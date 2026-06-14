from __future__ import annotations

import csv
import re
import shutil
import subprocess
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


REPO = Path(__file__).resolve().parents[1]
FINAL = REPO / "课程最终提交材料"
OUT_DIR = FINAL / "03_小组汇报PPT和报告"
DOCX = OUT_DIR / "企业AI部署偏好与治理机制研究_最终课程报告.docx"
PDF = OUT_DIR / "企业AI部署偏好与治理机制研究_最终课程报告.pdf"
TABLES = REPO / "outputs" / "tables"
FIGS = REPO / "outputs" / "figures" / "academic"
BLACK = RGBColor(0, 0, 0)


def read_csv(name: str) -> list[dict[str, str]]:
    with (TABLES / name).open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


ALG = read_csv("course_algorithm_comparison.csv")
REG = read_csv("course_regression_summary.csv")
VIF = read_csv("course_vif_diagnostics.csv")


def set_font(run, size: float = 10.5, bold: bool = False, font: str = "宋体") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = BLACK


def para(doc: Document, text: str = "", *, indent: bool = True, size: float = 10.5, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.45
    p.paragraph_format.space_after = Pt(5)
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold)


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, size=15 if level == 1 else 13, bold=True, font="黑体")


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-18)
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("• " + text)
    set_font(r, size=10.5)


def cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(text))
    set_font(r, size=9.5, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell_text(t.rows[0].cells[i], h, bold=True)
        shade(t.rows[0].cells[i], "D9D9D9")
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value)
    doc.add_paragraph()


def figure(doc: Document, filename: str, caption: str, width: float = 12.6) -> None:
    path = FIGS / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_font(r, size=9.5)


def normalize_docx_colors(path: Path) -> None:
    tmp = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                text = data.decode("utf-8", "ignore")
                text = re.sub(r'(<w:color\b[^>]*\bw:val=")[^"]+(")', r"\g<1>000000\2", text)
                data = text.encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(path)


def build() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.25)
    section.bottom_margin = Cm(2.1)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    styles = doc.styles
    for style_name in ["Normal", "Body Text"]:
        if style_name in styles:
            styles[style_name].font.name = "宋体"
            styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            styles[style_name].font.color.rgb = BLACK

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("机器学习课程最终小组案例报告")
    set_font(r, size=18, bold=True, font="黑体")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("企业AI部署偏好与治理机制研究")
    set_font(r, size=20, bold=True, font="黑体")
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run("基于Eurostat企业ICT数据的机器学习分析")
    set_font(r, size=14, bold=True, font="黑体")
    doc.add_paragraph()
    table(
        doc,
        ["项目", "内容"],
        [
            ["课程名称", "机器学习"],
            ["小组成员", "景浩伟、张新通、黄陈熙、刘子涵"],
            ["班级与学号", "信管2301｜景浩伟 202321054012"],
            ["提交结构", "01_数据、02_源码、03_小组汇报PPT和报告、04_个人作业总结"],
            ["仓库主题", "machine-learning-course-final-submission-2026"],
        ],
    )
    para(doc, "说明：本报告在原投稿稿研究内容基础上改写为课程作业版本，重点说明数据来源、代码流程、模型验证、结果解释、PPT证据映射和个人作业提交情况。报告中的指标来自本地仓库结果表，不把机器学习相关性写成因果证明。")
    doc.add_page_break()

    heading(doc, "摘要")
    for text in [
        "本课程案例围绕企业AI部署偏好与治理机制展开。项目没有采用虚构问卷或手工编造指标，而是以Eurostat官方企业ICT调查数据为主线，完成数据获取、清洗、特征工程、模型训练、分组交叉验证、图表生成和Agent原型说明。最终材料按数据、源码、小组汇报PPT和报告、个人作业总结四类组织，便于课程验收时逐项检查。",
        "研究部分分为两个阶段。Stage 1使用企业规模组数据解释中小企业AI流程自动化采纳机制，544个建模样本覆盖36个国家或地区；Stage 2使用行业和区域层面的GE10数据做外部验证，建模面板为5,814行，覆盖36个国家或地区和50个NACE行业。验证方法采用GroupKFold，以国家为分组变量，避免同一国家的相似观测同时进入训练集和测试集。",
        "结果显示，机器学习能力相关指标在两个阶段中都具有较高解释力；治理成熟度在行业区域层面更稳定；安全顾虑不宜简单理解为采纳阻力，更适合看作部署路径选择的影响因素。报告最后说明课程知识点对应关系、仓库整理边界和仍然存在的局限。",
    ]:
        para(doc, text)
    para(doc, "关键词：企业AI采纳；流程自动化；机器学习；GroupKFold；部署偏好；课程案例", indent=False, bold=True)

    doc.add_page_break()
    heading(doc, "一、课程提交材料整理说明")
    para(doc, "最终提交目录不是把所有历史文件简单打包，而是按老师检查课程作业时最自然的顺序重排。原始下载大文件、历史PPT多版本、临时预览图、payload、QA草稿和Python缓存不进入最终展示目录。保留下来的内容必须能回答四个问题：数据从哪里来，代码怎么跑，小组汇报如何展示，个人作业是否完整。")
    table(
        doc,
        ["顺序", "目录", "保留内容", "检查重点"],
        [
            ["01", "数据", "processed、samples、manifest、outputs表格/报告/图表", "数据来源、清洗结果、模型指标是否可追溯"],
            ["02", "源码", "src、configs、notebooks、10_Agent系统、requirements", "能否说明数据处理、训练和Agent原型"],
            ["03", "小组汇报PPT和报告", "最终PPT、PPT PDF、最终课程报告DOCX/PDF", "小组案例是否完整，报告是否超过15页"],
            ["04", "个人作业总结", "平时作业汇总、十次作业PDF、个人任务报告", "个人作业是否齐全，顺序是否清楚"],
        ],
    )

    doc.add_page_break()
    heading(doc, "二、案例选题与问题来源")
    for text in [
        "选题来自一个很具体的问题：企业已经开始尝试AI工具，但不同企业选择SaaS、API、本地化或混合部署的原因并不一样。单纯说企业愿不愿意用AI，解释力度不够；真正影响落地的是效率需求、安全顾虑、数字基础和治理能力之间的组合。",
        "课程案例将这个问题转成机器学习任务：以AI流程自动化使用率作为目标变量，用企业ICT能力、云计算能力、数据成熟度、安全顾虑和治理成熟度等指标做特征，训练可解释模型，观察哪些因素稳定影响采纳强度。这个任务既能对应回归、树模型、聚类和交叉验证，也能和小组PPT里的部署建议自然衔接。",
        "报告中的结论保持边界。Eurostat数据主要来自欧洲企业环境，Stage 2还是GE10行业区域口径，不能把Stage 2结果直接写成SME-only结论，也不能把非实验模型结果写成因果结论。",
    ]:
        para(doc, text)
    heading(doc, "2.1 理论框架：TOE与TAM的课程化改写", 2)
    para(doc, "投稿稿中使用技术—组织—环境（TOE）框架和技术接受模型（TAM）解释企业AI采纳。本课程报告保留这个理论骨架，但把它改成更便于机器学习课程检查的变量框架。TOE的技术维度对应AI能力、云服务、数据成熟度和数字基础；组织维度对应ICT人才、流程开发、培训和治理成熟度；环境维度对应国家、行业、年份和市场数字化差异。TAM中的感知有用性被转写为效率需求，感知易用性被转写为部署准备度。")
    para(doc, "这样处理的好处是，理论构念不再停留在文字描述，而能落到可计算指标上。比如，企业是否使用机器学习技术、是否使用云开发平台、是否具备数据分析能力，都可以从Eurostat企业ICT调查中找到对应指标。课程报告因此既能说明研究问题，也能说明模型特征从哪里来。")
    heading(doc, "2.2 安全顾虑不是单纯阻力", 2)
    para(doc, "投稿稿里一个重要判断是：安全顾虑不一定阻止企业采纳AI，而会改变企业的部署路径。这个判断在课程报告中继续保留。对于安全要求较低、流程标准化程度较高的企业，SaaS和API接入更容易落地；对于数据敏感、合规压力较高的企业，本地化、私有云或混合部署更合理。换句话说，安全顾虑更像部署路径转换器，而不是简单的负向变量。")

    doc.add_page_break()
    heading(doc, "三、数据来源与真实性核验")
    para(doc, "数据来自Eurostat官方SDMX接口，仓库保留manifest记录，包括源文件名称、URL、文件大小、下载状态和校验信息。最终展示版不上传全部raw大文件，是为了控制GitHub体积；但保留了manifest、README和下载脚本，老师需要复现时可以重新获取。")
    para(doc, "需要把数据量说准确：Stage 2官方源数据链是12,770,332行，也就是约1277万行，不是1.2亿行。项目并不是直接用千万级原始行训练模型，而是先做源文件剖析、指标筛选、覆盖率过滤、面板聚合和目标变量非空筛选，最后形成可解释、可复现的建模面板。这个过程本身就是数据挖掘课程里最重要的清洗环节。")
    table(
        doc,
        ["阶段", "数据用途", "样本口径", "建模规模", "说明"],
        [
            ["Stage 1", "SME规模组机制解释", "企业规模组、国家、年份", "544建模行", "用于解释中小企业相关机制"],
            ["Stage 2", "行业/区域外部验证", "GE10行业、国家、年份", "5,814建模行", "用于外部验证，不表述为SME-only"],
        ],
    )
    table(
        doc,
        ["数据链条", "数量"],
        [
            ["官方源文件记录行数", "12,770,332"],
            ["特征提取扫描行数", "12,341,630"],
            ["指标过滤后保留行数", "856,880"],
            ["最终Stage 2建模面板", "5,814"],
        ],
    )
    figure(doc, "fig1_academic_validation_clean.png", "图1 数据验证与建模流程图")
    para(doc, "清洗代码分布在src/acquisition、src/cleaning和src/pipeline相关脚本中。download_sources.py和download_stage2_large_sources.py负责官方数据下载与manifest登记；profile_stage2_sources.py负责源文件剖析；pipeline.py、pipeline_multisource.py和pipeline_stage2_large.py负责长表转换、变量筛选、面板合并、模型训练和结果登记。这些脚本共同说明了从源数据到建模数据的全过程。")

    doc.add_page_break()
    heading(doc, "四、数据清洗、特征工程与变量构建")
    for text in [
        "清洗的第一步是把不同Eurostat数据集整理成长表，再按国家、年份、企业规模或行业进行合并。由于不同指标覆盖年份和国家不完全一致，项目设置覆盖率阈值，剔除缺失过高或含义不稳定的字段。这样会减少变量数量，但后续模型更稳。",
        "特征工程围绕四类机制变量展开：效率需求、安全顾虑、部署准备度和治理成熟度。效率需求主要用AI技术能力、自然语言生成、认知计算等指标体现；部署准备度与云计算开发、云数据分析和数字基础有关；安全顾虑来自隐私、安全和数据管理相关指标；治理成熟度用ICT培训、流程开发和数据管理实践作为代理。",
        "因变量是AI流程自动化使用率。为了避免信息泄漏，模型训练时剔除了目标变量及其直接派生字段。线性模型用于看方向，树模型用于捕捉非线性关系，聚类用于把企业画像转成部署策略建议。",
    ]:
        para(doc, text)
    heading(doc, "4.1 清洗规则", 2)
    para(doc, "清洗规则主要包括四类。第一，统一不同Eurostat数据集中的国家、年份、行业和规模字段，避免同一含义出现多种编码。第二，把原始长表转为宽面板，使每一行对应一个国家—年份—规模组或国家—年份—行业观测。第三，剔除缺失率过高、覆盖口径不稳定或含义与目标变量过近的指标。第四，对模型训练特征做泄漏控制，删除目标变量、目标变量同义字段和显然由目标派生的gap字段。")
    para(doc, "这些清洗规则比简单删除空值更重要。企业ICT数据来自多年、多国、多行业调查，字段覆盖不均衡很正常。如果强行保留所有指标，模型会在大量缺失和不稳定口径中学习到噪声；如果只保留目标变量附近的指标，又容易产生信息泄漏。本项目在两者之间取一个课程上可解释的平衡。")
    heading(doc, "4.2 私密材料与公开材料边界", 2)
    para(doc, "投稿稿提到前期访谈和问卷材料，它们用于提出效率需求、安全顾虑和部署偏好三个构念，也用于解释模型结果。但这些材料包含企业具体情境，不作为公开仓库的逐行原始数据。公开仓库的主证据是Eurostat官方数据、清洗脚本、建模面板、模型结果表和图表。这种边界能同时满足课程复现、数据伦理和公开展示需要。")
    table(
        doc,
        ["变量类型", "代表变量或指标", "课程解释"],
        [
            ["因变量", "E_AI_TPA", "企业使用AI进行工作流自动化的比例"],
            ["效率需求", "E_AI_TML、E_AI_TNLG、E_AI_CC", "企业对AI提升流程效率的需求和能力基础"],
            ["部署准备度", "E_CC_PDEV、E_CC_DA、digital_foundation_index", "企业云服务、数据分析和数字基础"],
            ["治理成熟度", "governance_maturity_proxy", "企业是否具备规则、培训和数据管理基础"],
            ["安全顾虑", "security_concern_index", "企业对隐私、安全、合规的担心"],
            ["控制变量", "geo、year、nace_r2、size_emp", "国家、年份、行业和规模差异"],
        ],
    )

    doc.add_page_break()
    heading(doc, "五、源码结构与复现流程")
    para(doc, "源码放在最终提交目录的02_源码中，同时仓库根目录保留同样结构。src目录负责数据下载、清洗、建模和图表生成；10_Agent系统目录负责研究Agent原型；outputs目录保存运行结果。老师既能从报告看方法，也能从源码找到对应脚本。")
    table(
        doc,
        ["代码入口", "作用"],
        [
            ["src/acquisition/download_sources.py", "下载Stage 1官方数据并记录manifest"],
            ["src/acquisition/download_stage2_large_sources.py", "下载Stage 2行业/区域数据"],
            ["src/pipeline.py", "构建Stage 1面板并训练基础模型"],
            ["src/pipeline_multisource.py", "整合多源企业ICT数据"],
            ["src/pipeline_stage2_large.py", "构建Stage 2行业/区域面板并训练模型"],
            ["src/course_ml_diagnostics.py", "生成课程需要的OLS、VIF和模型比较"],
            ["src/render_academic_figures.py", "生成报告图表"],
            ["10_Agent系统/tests", "检查预测格式、引用准确性、泄漏控制和质量契约"],
        ],
    )
    para(doc, "整理后已在本地运行单元测试：python -m unittest discover -s \"10_Agent系统/tests\" -p \"test_*.py\"，结果为8项测试通过。由于最终展示版不上传大体积raw文件，测试契约检查manifest、processed面板和raw说明文件是否存在，而不是要求所有raw下载文件都在GitHub里。")
    para(doc, "复现时建议先看课程最终提交材料中的02_源码，再回到仓库根目录运行同名脚本。最终提交目录是给老师检查的精简版，仓库根目录是给复现实验使用的工作版。二者保留同样的核心脚本和结果文件，但最终提交目录去掉了历史PPT版本、临时预览图和不必要的大体积raw下载文件。")

    doc.add_page_break()
    heading(doc, "六、模型方法与训练设计")
    para(doc, "本案例使用的模型不追求复杂，而是追求课程知识点清楚。OLS用于看变量方向和显著性，Ridge用于处理线性模型中的共线性问题，Random Forest和ExtraTrees用于处理非线性关系，KMeans用于把企业分为不同画像。模型不是为了追求最高分，而是为了把课程算法放到一个完整数据项目中。")
    para(doc, "GroupKFold是本项目最关键的验证设计。企业ICT数据具有明显的国家和区域结构，同一国家的企业观测往往共享制度环境、数字基础和产业结构。如果随机划分训练集和测试集，同一国家的相似观测可能同时出现在两边，模型分数会被高估。按国家分组交叉验证会更严格，也更接近外部泛化场景。")
    para(doc, "Ridge回归适合处理多重共线性较明显的表格数据，ExtraTrees适合捕捉非线性和变量交互。A10 GPU上的神经网络基线被保留为训练记录，但不作为主要结论，因为结构化表格数据在样本量有限、特征解释要求较高时，树模型和正则化线性模型通常更稳。")
    table(
        doc,
        ["方法", "在本案例中的用途", "课程对应知识点"],
        [
            ["OLS", "查看变量方向、显著性和诊断指标", "线性回归、统计解释"],
            ["Ridge", "作为Stage 1公开基准模型", "正则化回归"],
            ["Random Forest / ExtraTrees", "处理非线性和变量交互", "决策树、随机森林"],
            ["GroupKFold", "按国家分组验证，降低地理泄漏", "训练集/测试集划分、交叉验证"],
            ["VIF", "检查多重共线性", "回归诊断"],
            ["KMeans", "识别企业画像并映射部署策略", "聚类算法"],
        ],
    )
    figure(doc, "fig1a_model_comparison_ppt.png", "图2 模型比较结果")

    doc.add_page_break()
    heading(doc, "七、主要结果与解释")
    heading(doc, "7.1 Stage 1：SME规模层机制解释", 2)
    s1 = [r for r in ALG if r["dataset"].startswith("Stage 1")]
    table(
        doc,
        ["模型", "R²均值", "MAE均值", "验证方式"],
        [[r["model_cn"], f"{float(r['r2_mean']):.4f}", f"{float(r['mae_mean']):.4f}", r["validation"]] for r in s1],
    )
    para(doc, "Stage 1使用544个建模样本，覆盖36个国家或地区。Ridge在GroupKFold by country下的R²均值为0.8744，MAE均值为1.7730。OLS结果显示，机器学习能力变量E_AI_TML的标准化系数最高，方向为正，说明已经具备机器学习技术基础的企业，更可能使用AI进行流程自动化。")
    para(doc, "投稿稿中还报告了完整训练集口径下Ridge拟合结果R²=0.8680、MAE=1.8342。课程诊断表中的GroupKFold结果与公开锁定指标口径不同，但方向一致：机器学习能力和云开发能力是Stage 1中最稳定的正向解释因素。云数据分析在OLS中出现负向系数，同时伴随较高VIF，这提醒我们解释单个系数时必须结合共线性诊断。")
    heading(doc, "7.2 Stage 2：行业/区域外部验证", 2)
    s2 = [r for r in ALG if r["dataset"].startswith("Stage 2")]
    table(
        doc,
        ["模型", "R²均值", "MAE均值", "验证方式"],
        [[r["model_cn"], f"{float(r['r2_mean']):.4f}", f"{float(r['mae_mean']):.4f}", r["validation"]] for r in s2],
    )
    para(doc, "Stage 2使用5,814行建模面板，覆盖36个国家或地区和50个NACE行业。ExtraTrees在10个核心特征口径下R²均值为0.7073，MAE均值为2.1060。换到行业和区域层面后，模型难度上升，但仍能保留较强解释力。")
    para(doc, "Stage 2的作用不是替代Stage 1，而是检验机制在更宽口径下是否还能成立。结果显示，机器学习能力仍然排名靠前，治理成熟度在行业/区域层面更清楚。这个差异符合直觉：单个规模组中治理指标可能被企业规模、年份和国家差异稀释，但在行业层面，治理能力对AI流程自动化的影响更容易体现出来。")
    heading(doc, "7.3 回归与共线性诊断", 2)
    table(
        doc,
        ["数据层", "n", "R²", "Adj.R²", "特征数", "p<0.05"],
        [[r["dataset"], r["n"], f"{float(r['r2']):.4f}", f"{float(r['adj_r2']):.4f}", r["features"], r["significant_05"]] for r in REG],
    )
    table(
        doc,
        ["数据层", "特征", "含义", "VIF"],
        [[r["dataset"], r["feature"], r["feature_label"], f"{float(r['vif']):.2f}"] for r in VIF[:8]],
    )
    para(doc, "VIF结果提醒我们：部分数字能力变量之间存在共线性，尤其是数据成熟度、部署准备度和云能力指标。报告解释时不把单个OLS系数当成唯一证据，而是结合Ridge、树模型特征重要性和外部验证结果判断。")
    para(doc, "从统计报告角度看，本项目已经报告样本量、R²、调整R²、MAE、VIF和GroupKFold验证方式。它没有做随机实验，也没有使用企业级干预数据，因此所有结论都表述为稳定关联、预测机制和部署启示，而不是因果效应。")
    figure(doc, "fig1b_sme_importance_ppt.png", "图3 Stage 1特征重要性")
    figure(doc, "fig2_stage2_external_importance.png", "图4 Stage 2外部验证特征重要性")

    doc.add_page_break()
    heading(doc, "八、部署偏好与企业画像")
    para(doc, "安全敏感型企业并不是完全拒绝AI，而是更适合本地化、混合云、审计日志和权限分级等部署方式。效率需求强但治理基础一般的企业，更适合从标准SaaS或API接入开始；治理基础成熟的企业，可以推进流程级自动化与跨系统集成。")
    bullet(doc, "高采纳型：数字基础和治理能力较好，可以推进流程级AI自动化。")
    bullet(doc, "稳健增长型：已有云服务和数据分析基础，适合从局部流程扩展到关键业务流程。")
    bullet(doc, "安全敏感型：优先考虑本地化、私有云、混合部署、日志审计和权限分级。")
    bullet(doc, "初始观望型：先做低风险SaaS试点，积累数据治理经验后再扩大应用范围。")
    para(doc, "这部分是小组汇报最容易讲清楚的业务转化。模型告诉我们哪些因素和AI流程自动化采纳更稳定相关，聚类再把企业分成可理解画像。展示时不要只讲模型分数，而要讲模型结果如何落到企业部署策略：谁适合轻量试点，谁需要治理先行，谁需要安全架构，谁可以推进流程级集成。")

    doc.add_page_break()
    heading(doc, "九、图表、PPT与报告之间的对应")
    para(doc, "小组汇报PPT不是单独做出来的展示稿，而是从数据、模型和报告结果中提炼出来的。PPT中的数据生命周期、模型验证、特征重要性、Agent质量评估等页面，都能在outputs目录和本报告中找到对应证据。最终目录保留了PPT、PPT PDF、contact sheet和slide_to_evidence_map。")
    para(doc, "图表专业度主要从三方面控制。第一，图表只展示关键证据，不使用无关装饰背景。第二，PNG图表分辨率满足课堂投影和PDF查看，核心图表宽度约1400到2100像素。第三，每张PPT页都有证据映射表，能说明使用了哪个结果表、哪张图或哪个报告段落。")
    table(
        doc,
        ["展示材料", "文件", "对应证据"],
        [
            ["小组汇报PPT", "中小企业AI流程自动化采纳机制研究案例_小组最终汇报PPT.pptx", "PPT可编辑原件"],
            ["PPT导出PDF", "中小企业AI流程自动化采纳机制研究案例_小组最终汇报PPT.pdf", "便于直接打开审核"],
            ["证据映射表", "slide_to_evidence_map_小组最终汇报PPT.csv", "说明每页PPT使用的数据来源"],
            ["预览总览图", "contact_sheet_小组最终汇报PPT.png", "快速检查版式和页序"],
        ],
    )

    doc.add_page_break()
    heading(doc, "十、Agent原型与质量检查")
    para(doc, "10_Agent系统是本项目的扩展部分。它的目的不是把报告变成聊天机器人，而是把数据查询、模型预测、证据引用和回答约束做成可测试的工具。query_indicator.py负责查指标，predict_adoption.py负责预测，cite_source.py负责证据引用，agent_answer.py负责把工具结果组织成回答。")
    para(doc, "Agent原型的质量标准不是回答越长越好，而是回答必须受证据约束。没有证据的问题返回无法确认，引用必须来自仓库中的结果表、报告或数据说明。这一点和课程作业的完整性有关：模型、报告、PPT和Agent都应指向同一套证据，而不是各自生成一套说法。")
    table(
        doc,
        ["测试项", "检查内容"],
        [
            ["prediction schema", "预测输出字段是否符合约定"],
            ["citation accuracy", "引用是否来自仓库证据"],
            ["no data leakage", "训练特征是否包含目标泄漏字段"],
            ["agent quality contract", "无证据问题是否返回无法确认，RAG数量是否受控"],
        ],
    )

    doc.add_page_break()
    heading(doc, "十一、课程知识点对应与个人作业说明")
    para(doc, "从课程角度看，本案例覆盖了数据挖掘流程、回归模型、树模型思想、聚类分析、特征工程、训练测试划分、指标解释和结果可视化。平时作业中做过的线性回归、Logistic/KNN、决策树、随机森林、KMeans等内容，在这个案例中都有对应位置。")
    table(
        doc,
        ["课程内容", "本案例对应"],
        [
            ["数据挖掘流程", "数据来源、清洗、特征工程、建模、解释、展示"],
            ["线性回归/多元回归", "OLS与Ridge模型"],
            ["决策树/随机森林", "Random Forest和ExtraTrees模型对比"],
            ["聚类分析", "KMeans企业画像和部署策略映射"],
            ["训练集与测试集", "GroupKFold按国家分组验证"],
            ["模型评价指标", "R²、MAE、VIF、特征重要性"],
            ["可视化", "报告图表、PPT图表和contact sheet"],
        ],
    )
    para(doc, "个人作业部分已经放在04_个人作业总结中，包括十次作业PDF提交版、平时作业汇总PDF和个人任务报告。十次作业PDF按01到10排序，便于逐份检查。代码类作业保留了代码、结果、图表或混淆矩阵等内容，不再提交PPTX原件，避免材料混乱。")

    doc.add_page_break()
    heading(doc, "十二、局限、改进方向与结论")
    for text in [
        "本项目仍有局限。第一，Stage 2使用GE10口径，不能直接等同于中小企业样本。第二，部分变量存在多重共线性，尤其在Stage 1中需要谨慎解释单个OLS系数。第三，Eurostat数据来自欧洲企业环境，不能直接代表中国中小企业平均情况。第四，访谈和问卷材料只用于构念提出和解释辅助，没有把逐行原始数据放入公开仓库。",
        "后续如果继续完善，可以加入中国本土企业ICT调查数据，或者把企业级问卷、访谈材料和公开统计数据做更细的对照。模型层面可以进一步尝试时间留出验证和行业留出验证。应用层面可以把Agent原型接入真实业务流程，观察它在文档处理、知识检索和客户服务场景中的实际效果。",
        "总的来说，本案例完成了课程要求中的小组案例数据、代码、PPT和15页以上报告，也把个人平时作业和个人总结按顺序放入最终提交目录。相比原来的报告，新版报告更强调课程过程和可检查材料，避免把展示词写得过满。机器学习在这里的作用是帮助我们看清数据中的稳定关系，而不是替代业务判断。",
    ]:
        para(doc, text)

    doc.add_page_break()
    heading(doc, "附录A：提交文件清单")
    table(
        doc,
        ["目录", "主要文件"],
        [
            ["01_数据", "processed、samples、raw_manifests、outputs_tables、outputs_reports、outputs_figures、数据来源说明"],
            ["02_源码", "src、configs、notebooks、10_Agent系统、requirements.txt、RUN_ALL_CHECKS.md"],
            ["03_小组汇报PPT和报告", "最终PPT、PPT PDF、证据映射、contact sheet、最终课程报告DOCX/PDF"],
            ["04_个人作业总结", "平时作业汇总、个人任务报告、十次作业PDF提交版"],
        ],
    )
    para(doc, "最终提交目录没有保留历史PPT多版本、过程预览图、构建payload、Python缓存和大体积raw下载文件。需要复现raw数据时，可通过src/acquisition中的下载脚本重新获取。")

    heading(doc, "附录B：参考资料")
    for ref in [
        "Eurostat Data Browser and SDMX API：企业ICT、AI、云计算、数字强度等公开统计数据。",
        "NIST AI Risk Management Framework 1.0：用于理解AI治理、安全和风险管理。",
        "scikit-learn documentation：GroupKFold、Ridge、RandomForest、ExtraTrees、KMeans等模型方法。",
        "投稿稿_科技管理研究_终版_改写.docx：本课程报告的研究内容基础，已改写为课程作业版。",
    ]:
        para(doc, ref, indent=False)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("机器学习课程最终小组案例报告｜景浩伟 202321054012")
    set_font(r, size=9)

    doc.save(DOCX)
    normalize_docx_colors(DOCX)

    if PDF.exists():
        PDF.unlink()
    subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(OUT_DIR), str(DOCX)],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    generated = OUT_DIR / (DOCX.stem + ".pdf")
    if generated != PDF and generated.exists():
        shutil.move(str(generated), str(PDF))


if __name__ == "__main__":
    build()
    print(DOCX)
