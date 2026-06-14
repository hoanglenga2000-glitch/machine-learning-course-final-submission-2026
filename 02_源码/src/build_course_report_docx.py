from __future__ import annotations

import json
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
PKG = ROOT / ".python_packages"
if PKG.exists() and str(PKG) not in sys.path:
    sys.path.insert(0, str(PKG))

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


REPO = ROOT
TABLES = REPO / "outputs" / "tables"
REPORTS = REPO / "outputs" / "reports"
FIGS = REPO / "05_学术图表" / "汇报图片稿_4K待审核" / "图表"
OUT_DIR = REPO / "06_结课报告"
OUT_DIR.mkdir(parents=True, exist_ok=True)
DOCX = OUT_DIR / "中小企业AI流程自动化采纳机制研究_机器学习结课报告.docx"


def csv(name: str) -> pd.DataFrame:
    return pd.read_csv(TABLES / name)


quality = csv("enhanced_data_quality_audit.csv")
cv = csv("enhanced_cv_results.csv")
gpu = csv("enhanced_gpu_baseline.csv")
alg = csv("course_algorithm_comparison.csv")
reg = csv("course_regression_summary.csv")
ols = csv("course_ols_coefficients.csv")
perm = csv("enhanced_permutation_importance.csv")


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color="111827", size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(t.rows[0].cells[i], h, bold=True, color="FFFFFF", size=8.5)
        set_cell_shading(t.rows[0].cells[i], "08224A")
        if widths:
            t.columns[i].width = Cm(widths[i])
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            set_cell_text(cells[i], v, size=8.3)
            set_cell_shading(cells[i], "F8FBFF" if len(t.rows) % 2 else "FFFFFF")
    doc.add_paragraph()
    return t


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.color.rgb = RGBColor(8, 34, 74)
    return p


def add_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(99, 112, 131)


def add_figure(doc, filename, caption):
    path = FIGS / filename
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(6.2))
        add_caption(doc, caption)


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.2)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)

styles = doc.styles
styles["Normal"].font.name = "Microsoft YaHei"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
styles["Normal"].font.size = Pt(10.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("基于中小企业 AI 流程自动化采纳机制研究")
r.bold = True
r.font.name = "Microsoft YaHei"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(8, 34, 74)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("效率需求、安全顾虑与部署偏好的机器学习实证分析")
r.font.name = "Microsoft YaHei"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(20, 100, 165)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("机器学习/数据挖掘课程结题案例报告｜数据来源：Eurostat 官方 SDMX-CSV｜应用场景：ai.zhjjq.tech")
r.font.name = "Microsoft YaHei"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
r.font.size = Pt(9.5)
r.font.color.rgb = RGBColor(99, 112, 131)

add_heading(doc, "摘要", 1)
add_para(doc, "本研究围绕中小企业 AI 流程自动化采纳机制，使用 Eurostat 官方企业数字化与 AI 采用数据，构建从源数据采集、清洗审计、特征工程、机器学习训练到部署策略建议的完整数据挖掘案例。研究将效率需求、安全顾虑与部署准备度转化为可计算变量，采用多元线性回归、岭回归、随机森林、ExtraTrees 和 A10 GPU PyTorch MLP 进行对照训练。结果显示，机器学习能力、数字基础、治理成熟度和部署准备度是 AI 流程自动化采纳强度的重要解释因素；安全顾虑并非单纯抑制因素，而是影响企业选择 SaaS、API、本地化或混合部署的重要分流机制。")

add_heading(doc, "一、数据来源与真实性", 1)
add_para(doc, "最终进入训练的数据来自 Eurostat 官方 SDMX-CSV 接口。每个原始文件均记录 URL、HTTP 状态、下载时间、文件大小和 SHA256 校验值，保存在 data/raw/manifest.jsonl 与 data/raw/manifest_stage2.jsonl。仓库同时提供中文目录 01_源数据，便于课程检查。")
add_bullet(doc, "Stage 1 使用企业规模组口径，覆盖 AI 采用、云服务、数字强度、数据分析、ICT 人才与电子商务等官方指标。")
add_bullet(doc, "Stage 2 使用 GE10 企业规模口径下的 NACE 行业/区域数据，用作外部验证，不表述为 SME 规模拆分。")
add_bullet(doc, "曾尝试获取 U.S. Census BTOS AI 文件但接口返回 HTTP 403，未进入训练，也不作为结论来源。")

add_heading(doc, "二、数据生命周期与清洗结果", 1)
table(
    doc,
    ["阶段", "数量/结果"],
    [
        ["Stage 2 官方源文件", "17"],
        ["原始源数据扫描行数", "12,770,332"],
        ["非空观测值", "10,453,354"],
        ["机制变量筛选后保留行", "856,880"],
        ["Stage 2 行业/区域建模面板", "5,814 行 × 80 列"],
        ["Stage 1 SME 规模层建模样本", "544 行 × 100 列"],
    ],
    [6, 8],
)
add_figure(doc, "图01_数据生命周期漏斗.png", "图1 数据生命周期：从官方源数据到可训练面板")
add_para(doc, "清洗过程中执行了长表转面板、目标变量非空筛选、重复键检查、缺失率审计和目标泄漏控制。直接目标变量、总体 AI 采用变量以及 target/workflow_gap/adoption_gap 等目标派生变量不进入特征集。")

add_heading(doc, "三、变量体系与研究框架", 1)
add_para(doc, "本研究将 TOE 与 TAM 框架转化为可训练变量。技术条件包括机器学习能力、自然语言生成、云开发能力和数据分析成熟度；组织基础包括数字基础、ICT 人才、治理成熟度和部署准备度；环境压力包括市场数字化、行业差异、国家/区域异质性和时间趋势；感知机制包括效率需求、安全顾虑与采纳强度转化。")

add_heading(doc, "四、机器学习方法", 1)
add_para(doc, "课程案例采用监督学习与解释模型结合的方式：OLS 用于解释机制方向与显著性；岭回归处理多重共线性；随机森林与 ExtraTrees 捕捉非线性和变量交互；A10 GPU 上的 PyTorch MLP 作为神经网络基线。模型评估采用 R2 和 MAE，并使用按国家分组的 GroupKFold 交叉验证，避免同一国家观测同时进入训练集和测试集造成泛化能力虚高。")
add_figure(doc, "图02_模型交叉验证比较.png", "图2 模型交叉验证比较")

add_heading(doc, "五、回归解释与特征重要性", 1)
reg_rows = []
for _, row in reg.iterrows():
    reg_rows.append([row["dataset"], int(row["n"]), f"{row['r2']:.3f}", f"{row['adj_r2']:.3f}", int(row["features"]), int(row["significant_05"])])
table(doc, ["模型层级", "n", "R2", "Adj.R2", "特征数", "p<0.05"], reg_rows, [4, 2, 2, 2, 2, 2])
add_figure(doc, "图03_多元回归标准化系数.png", "图3 多元回归标准化系数")
add_figure(doc, "图04_特征重要性双面板.png", "图4 随机森林/ExtraTrees 特征重要性")
add_para(doc, "回归与特征重要性结果共同显示，机器学习能力是最稳定的核心驱动。Stage 2 中数字基础、治理成熟度和部署准备度也呈现显著正向解释力，说明 AI 流程自动化采纳不是单一工具购买行为，而是技术、组织与治理能力共同作用的结果。")

add_heading(doc, "六、A10 GPU 基线与模型选择", 1)
gpu_rows = []
for _, row in gpu.iterrows():
    gpu_rows.append([
        "SME规模层" if row["dataset"] == "stage1_sme_size_class" else "行业验证层",
        row["gpu_name"],
        row["gpu_model"],
        f"{row['r2']:.3f}",
        f"{row['mae']:.3f}",
        int(row["best_epoch"]),
    ])
table(doc, ["数据层", "设备", "模型", "R2", "MAE", "最佳epoch"], gpu_rows, [3, 3, 3, 2, 2, 2])
add_figure(doc, "图05_A10_GPU_MLP基线.png", "图5 A10 GPU PyTorch MLP 基线")
add_para(doc, "A10 GPU 基线结果低于树模型，说明结构化企业面板数据并不一定适合更复杂的神经网络模型。该结果体现了机器学习课程中的重要原则：模型复杂度必须服从数据结构、样本量与泛化验证结果。")

add_heading(doc, "七、部署策略与企业价值", 1)
add_figure(doc, "图07_部署偏好策略矩阵.png", "图6 部署偏好策略矩阵")
add_para(doc, "安全顾虑不是简单阻碍 AI 采用，而是企业部署偏好的分流器。高安全顾虑企业更适合本地化或私有云部署；高效率需求但安全要求较高的企业适合混合部署；低安全顾虑且流程集成需求高的企业适合 API 接入；基础薄弱企业更适合标准 SaaS 试点。该结论可直接服务 ai.zhjjq.tech 的客户分层、流程入口、审批治理和部署报价策略。")

add_heading(doc, "八、结论与后续研究", 1)
add_bullet(doc, "结论一：企业 AI 流程自动化采纳强度主要由机器学习能力、数据基础、部署准备度和治理成熟度共同驱动。")
add_bullet(doc, "结论二：安全顾虑会重塑部署路径，使企业从单纯公有 SaaS 转向本地化、私有云或混合部署。")
add_bullet(doc, "结论三：SME 层随机森林 R2=0.850，行业/区域外部验证层 ExtraTrees R2=0.724，机制具有一定迁移性。")
add_bullet(doc, "结论四：A10 GPU MLP 低于树模型，证明本项目更适合表格学习模型，而非盲目堆叠复杂神经网络。")
add_para(doc, "后续研究可接入国内中小企业真实使用日志、访谈和问卷追踪，结合 ai.zhjjq.tech 后台任务日志进行 A/B 测试，并加入面板固定效应或因果推断方法，进一步提升研究的解释深度和产品落地价值。")

add_heading(doc, "附录：提交目录说明", 1)
table(
    doc,
    ["中文目录", "内容"],
    [
        ["01_源数据", "Eurostat 官方原始下载数据、压缩文件、manifest 来源校验。"],
        ["02_清洗后数据", "清洗后的建模面板和样本预览。"],
        ["03_清洗与训练代码", "数据下载、清洗、特征工程、A10训练、课程诊断、图表生成代码。"],
        ["04_分析结果表格", "数据质量审计、模型指标、OLS、VIF、特征重要性、GPU基线。"],
        ["05_学术图表", "PNG/SVG 学术图表。"],
        ["06_结课报告", "数据来源、训练核验、研究质量说明和最终报告。"],
    ],
    [4, 10],
)

doc.save(DOCX)
print(DOCX)
